from docx import Document
from docx.shared import Inches, Cm

# from docxtpl import DocxTemplate
# import win32com.client
# import inspect, os
# import pythoncom

report_data_path = "./main/static/report/"

# report_data_path = ''

# def update_toc(docx_file):
#     pythoncom.CoInitialize()
#     word = win32com.client.DispatchEx("Word.Application")
#     pythoncom.CoInitialize()
#     doc = word.Documents.Open(docx_file)
#     doc.TablesOfContents(1).Update()
#     doc.Close(SaveChanges=True)
#     word.Quit()

def renew_text(p, text):
    if "“" in text or "增速=" in text:
        return
    p.clear()
    p_text = p.add_run(text)
    p_text.font.name = 'Times New Roman'
    '''
    if "增速=" in text:
        print(p)
        print(text)

    p.clear()
    p.add_run()
    p.runs[0].text = text
    p.runs[0].font.name = 'Times New Roman'
    '''


def replace_pic(p, pic_list, width=None, height=None):
    p.clear()
    p.add_run()
    for pic in pic_list:
        pic = report_data_path + pic
        if width:
            p.runs[0].add_picture(pic, width=width)
        elif height:
            p.runs[0].add_picture(pic, height=height)


def createReport(report_name, num=33, compress_threshold=33):
    document = Document(report_data_path + '模板.docx')
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    paragraphs = document.paragraphs
    i = 0
    for para in paragraphs:
        print(i)
        print(para.text)
        images = para._element.xpath('.//pic:pic')
        for im in images:
            print(im)
        i += 1

    texts = []
    with open(report_data_path + 'text.txt', encoding='utf-8') as f:
        for line in f:
            texts.append(line.strip())
    to_replace_text_paragraphs = [
        paragraphs[4],
        paragraphs[6],
        paragraphs[20],
        paragraphs[37],
        paragraphs[38],
        paragraphs[39],
        paragraphs[40],
        paragraphs[58],
        paragraphs[59],
        paragraphs[60],
        paragraphs[61],
        paragraphs[63],
        paragraphs[65],
        paragraphs[68],
        paragraphs[71],
        paragraphs[72],
        paragraphs[76],
        paragraphs[78],
        paragraphs[80],
        paragraphs[86],
        paragraphs[88],
        paragraphs[90],
        paragraphs[93],
        paragraphs[96],
        paragraphs[99],
        paragraphs[101],
        paragraphs[104],
        paragraphs[105],
        paragraphs[107],
        paragraphs[108],
        paragraphs[110],
        paragraphs[111],
        paragraphs[113],
        paragraphs[114],
        paragraphs[116],
        paragraphs[117],
        paragraphs[123],
        paragraphs[125],
        paragraphs[127],
        paragraphs[129],
        paragraphs[131],
        paragraphs[133],
        paragraphs[135],
        paragraphs[137],
        paragraphs[139],
        paragraphs[141]
    ]
    if len(to_replace_text_paragraphs) != len(texts):
        print('error: inconsistent texts')
        print(len(to_replace_text_paragraphs))
        print(len(texts))
    else:
        for i in range(0, len(texts)):
            if i > 0:
                renew_text(to_replace_text_paragraphs[i], texts[i])
    # 替换文本，需要使用renew_text的方法，有三个参数分别是文档对应段落以及新的文本
    # renew_text(paragraphs[6],week_period)
    # renew_text(paragraphs[20],report_date)
    # renew_text(paragraphs[33],title_2_1)
    # renew_text(paragraphs[34],major_distribution)
    # renew_text(paragraphs[50],title_2_2)
    # renew_text(paragraphs[59],title_2_3)
    # renew_text(paragraphs[62],title_2_4)
    # renew_text(paragraphs[65],title_2_5)
    # renew_text(paragraphs[68],title_2_6)

    # 替换图片，需要使用replace_pic方法，有三个参数分别是文档对应段落、图片地址以及图片宽度

    pic_2_1 = ['2_1.png']
    pic_2_2 = ['2_2.png']
    pic_2_3 = ['2_3_a.png', '2_3_b.png']
    pic_2_4 = ['2_4_a.png', '2_4_b.png']
    pic_2_5 = ['2_5_a.png', '2_5_b.png']
    pic_2_6 = ['2_6.png']
    pic_2_7 = ['2_7.png']
    pic_2_8 = ['2_8.png']
    pic_2_9 = ['2_9.png']
    pic_2_10 = ['2_10.png']
    pic_2_11 = ['2_11.png']
    pic_2_12 = ['2_12.png']
    pic_2_13 = ['2_13_a.png', '2_13_b.png']
    pic_2_14 = ['2_14.png']
    pic_2_15 = ['2_15.png']
    pic_2_16 = ['2_16.png']
    pic_2_17 = ['2_17.png']
    pic_2_18 = ['2_18.png']
    pic_2_19 = ['2_19.png']
    pic_2_20 = ['2_20.png']
    pic_2_21 = ['2_21.png']
    pic_2_22 = ['2_22.png']
    pic_2_23 = ['2_23.png']
    pic_2_24 = ['2_24.png']
    pic_2_25 = ['2_25.png']
    pic_2_26 = ['2_26.png']

    replace_pic(paragraphs[62], pic_2_1, height=Cm(8.6))
    replace_pic(paragraphs[64], pic_2_2, height=Cm(8.6))
    replace_pic(paragraphs[66], pic_2_3, width=Cm(7.3))
    replace_pic(paragraphs[69], pic_2_4, width=Cm(7.3))
    replace_pic(paragraphs[74], pic_2_5, width=Cm(7.3))
    replace_pic(paragraphs[77], pic_2_6, height=Cm(8.28))
    replace_pic(paragraphs[79], pic_2_7, height=Cm(8.28))
    replace_pic(paragraphs[87], pic_2_8, height=Cm(21.51))
    replace_pic(paragraphs[89], pic_2_9, height=Cm(21.51))
    replace_pic(paragraphs[92], pic_2_10, height=Cm(21.51))
    replace_pic(paragraphs[95], pic_2_11, height=Cm(21.51))
    replace_pic(paragraphs[98], pic_2_12, height=Cm(21.51))
    replace_pic(paragraphs[102], pic_2_13, width=Cm(6.96))
    replace_pic(paragraphs[106], pic_2_14, height=Cm(5.7))
    replace_pic(paragraphs[109], pic_2_15, height=Cm(5.7))
    replace_pic(paragraphs[112], pic_2_16, width=Cm(10.9 if num <= compress_threshold else 14.5))
    replace_pic(paragraphs[115], pic_2_17, width=Cm(10.9 if num <= compress_threshold else 14.5))
    replace_pic(paragraphs[124], pic_2_18, height=Cm(21.27))
    replace_pic(paragraphs[126], pic_2_19, height=Cm(10))
    replace_pic(paragraphs[128], pic_2_20, height=Cm(10))
    replace_pic(paragraphs[130], pic_2_21, height=Cm(10))
    replace_pic(paragraphs[132], pic_2_22, height=Cm(10))
    replace_pic(paragraphs[134], pic_2_23, height=Cm(10))
    replace_pic(paragraphs[136], pic_2_24, height=Cm(10))
    replace_pic(paragraphs[138], pic_2_25, height=Cm(10))
    replace_pic(paragraphs[140], pic_2_26, height=Cm(10))

    file_name = report_data_path + report_name

    # script_dir = os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))
    # file_path = os.path.join(script_dir, file_name)

    # print(file_path)
    document.save(file_name)  # 保存文档
    # update_toc(file_path)


if __name__ == "__main__":
    # 33 10.7
    createReport()
